import random
import sys
import math
from functools import reduce
import docx
from functools import partial
from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
from PyQt5.QtCore import *
import sqlite3
from PyQt5 import QtCore, QtGui, QtWidgets
import datetime
from config import login, password
import barcode
from math import pi, cos, sin

current_password = password


class MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(799, 781)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.design_up = QtWidgets.QFrame(self.centralwidget)
        self.design_up.setGeometry(QtCore.QRect(0, 70, 801, 16))
        self.design_up.setStyleSheet("background-color: rgb(0, 170, 255);")
        self.design_up.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.design_up.setFrameShadow(QtWidgets.QFrame.Raised)
        self.design_up.setObjectName("design_up")
        self.phone = QtWidgets.QLabel(self.centralwidget)
        self.phone.setGeometry(QtCore.QRect(590, 50, 201, 16))
        self.phone.setObjectName("phone")
        self.lable_scan = QtWidgets.QLabel(self.centralwidget)
        self.lable_scan.setGeometry(QtCore.QRect(10, 100, 751, 41))
        font = QtGui.QFont()
        font.setPointSize(20)
        self.lable_scan.setFont(font)
        self.lable_scan.setAlignment(QtCore.Qt.AlignCenter)
        self.lable_scan.setObjectName("lable_scan")
        self.shopping = QtWidgets.QFrame(self.centralwidget)
        self.shopping.setGeometry(QtCore.QRect(20, 130, 391, 601))
        self.shopping.setStyleSheet("")
        self.shopping.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.shopping.setFrameShadow(QtWidgets.QFrame.Raised)
        self.shopping.setObjectName("shopping")
        self.purchase = QtWidgets.QLabel(self.shopping)
        self.purchase.setGeometry(QtCore.QRect(0, 10, 321, 31))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setBold(True)
        font.setWeight(75)
        self.purchase.setFont(font)
        self.purchase.setObjectName("purchase")
        self.col_items = QtWidgets.QTableWidget(self.shopping)
        self.col_items.setGeometry(QtCore.QRect(0, 50, 391, 481))
        self.col_items.setObjectName("col_items")
        self.col_items.setColumnCount(0)
        self.col_items.setRowCount(0)
        self.to_paid = QtWidgets.QLabel(self.shopping)
        self.to_paid.setGeometry(QtCore.QRect(0, 540, 161, 31))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setBold(True)
        font.setWeight(75)
        self.to_paid.setFont(font)
        self.to_paid.setObjectName("to_paid")
        self.money = QtWidgets.QTextBrowser(self.shopping)
        self.money.setGeometry(QtCore.QRect(170, 540, 221, 31))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.money.setFont(font)
        self.money.setObjectName("money")
        self.time = QtWidgets.QLabel(self.centralwidget)
        self.time.setGeometry(QtCore.QRect(30, 40, 171, 16))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.time.setFont(font)
        self.time.setObjectName("time")
        self.designdown = QtWidgets.QFrame(self.centralwidget)
        self.designdown.setGeometry(QtCore.QRect(-10, 730, 801, 21))
        self.designdown.setStyleSheet("background-color: rgb(0, 170, 255);")
        self.designdown.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.designdown.setFrameShadow(QtWidgets.QFrame.Raised)
        self.designdown.setObjectName("designdown")
        self.label = QtWidgets.QLabel(self.designdown)
        self.label.setGeometry(QtCore.QRect(20, 0, 771, 20))
        font = QtGui.QFont()
        font.setPointSize(8)
        self.label.setFont(font)
        self.label.setStyleSheet("color: rgb(170, 0, 0);")
        self.label.setAlignment(QtCore.Qt.AlignCenter)
        self.label.setObjectName("label")
        self.image_scan = QtWidgets.QLabel(self.centralwidget)
        self.image_scan.setGeometry(QtCore.QRect(430, 180, 351, 251))
        self.image_scan.setText("")
        self.image_scan.setPixmap(QtGui.QPixmap("icons/scan.jpg"))
        self.image_scan.setScaledContents(True)
        self.image_scan.setObjectName("image_scan")
        self.logo_auto = QtWidgets.QPushButton(self.centralwidget)
        self.logo_auto.setGeometry(QtCore.QRect(10, 10, 781, 41))
        font = QtGui.QFont()
        font.setFamily("Segoe Script")
        font.setPointSize(14)
        self.logo_auto.setFont(font)
        self.logo_auto.setStyleSheet("border: none;")
        self.logo_auto.setObjectName("logo_auto")
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setGeometry(QtCore.QRect(410, 760, 371, 20))
        font = QtGui.QFont()
        font.setPointSize(8)
        self.label_2.setFont(font)
        self.label_2.setStyleSheet("color: rgb(170, 0, 0);")
        self.label_2.setObjectName("label_2")
        self.start_shopping = QtWidgets.QPushButton(self.centralwidget)
        self.start_shopping.setGeometry(QtCore.QRect(610, 580, 171, 121))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.start_shopping.setFont(font)
        self.start_shopping.setStyleSheet("background-color: rgb(0, 170, 255);")
        self.start_shopping.setObjectName("start_shopping")
        self.small_package = QtWidgets.QPushButton(self.centralwidget)
        self.small_package.setGeometry(QtCore.QRect(430, 440, 171, 131))
        self.small_package.setStyleSheet("background-color: rgb(165, 165, 165);")
        self.small_package.setText("")
        self.small_package.setObjectName("small_package")
        self.label_10 = QtWidgets.QLabel(self.centralwidget)
        self.label_10.setGeometry(QtCore.QRect(460, 450, 111, 16))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_10.setFont(font)
        self.label_10.setObjectName("label_10")
        self.label_11 = QtWidgets.QLabel(self.centralwidget)
        self.label_11.setGeometry(QtCore.QRect(470, 470, 91, 81))
        self.label_11.setText("")
        self.label_11.setPixmap(QtGui.QPixmap("icons/pngwing.com.png"))
        self.label_11.setScaledContents(True)
        self.label_11.setAlignment(QtCore.Qt.AlignCenter)
        self.label_11.setObjectName("label_11")
        self.big_package = QtWidgets.QPushButton(self.centralwidget)
        self.big_package.setGeometry(QtCore.QRect(610, 440, 171, 131))
        self.big_package.setStyleSheet("background-color: rgb(165, 165, 165);")
        self.big_package.setText("")
        self.big_package.setObjectName("big_package")
        self.label_13 = QtWidgets.QLabel(self.centralwidget)
        self.label_13.setGeometry(QtCore.QRect(660, 470, 101, 91))
        self.label_13.setText("")
        self.label_13.setPixmap(QtGui.QPixmap("icons/pngwing.com (1).png"))
        self.label_13.setScaledContents(True)
        self.label_13.setAlignment(QtCore.Qt.AlignCenter)
        self.label_13.setObjectName("label_13")
        self.label_12 = QtWidgets.QLabel(self.centralwidget)
        self.label_12.setGeometry(QtCore.QRect(640, 450, 111, 16))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_12.setFont(font)
        self.label_12.setObjectName("label_12")
        self.manual_input = QtWidgets.QPushButton(self.centralwidget)
        self.manual_input.setGeometry(QtCore.QRect(430, 580, 171, 121))
        self.manual_input.setStyleSheet("background-color: rgb(165, 165, 165);")
        self.manual_input.setText("")
        self.manual_input.setObjectName("manual_input")
        self.label_7 = QtWidgets.QLabel(self.centralwidget)
        self.label_7.setGeometry(QtCore.QRect(450, 590, 131, 16))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_7.setFont(font)
        self.label_7.setAlignment(QtCore.Qt.AlignCenter)
        self.label_7.setObjectName("label_7")
        self.label_9 = QtWidgets.QLabel(self.centralwidget)
        self.label_9.setGeometry(QtCore.QRect(450, 620, 131, 61))
        self.label_9.setText("")
        self.label_9.setPixmap(QtGui.QPixmap("icons/trich.jpg"))
        self.label_9.setScaledContents(True)
        self.label_9.setAlignment(QtCore.Qt.AlignCenter)
        self.label_9.setObjectName("label_9")
        MainWindow.setCentralWidget(self.centralwidget)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.phone.setText(_translate("MainWindow", "телефон поддержки: 7-928-XXX-XX-XX"))
        self.lable_scan.setText(_translate("MainWindow", "Просканируйте товар"))
        self.purchase.setText(_translate("MainWindow", "Ваша покупка:"))
        self.to_paid.setText(_translate("MainWindow", "К оплате"))
        self.money.setHtml(_translate("MainWindow",
                                      "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
                                      "<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
                                      "p, li { white-space: pre-wrap; }\n"
                                      "</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:12pt; font-weight:400; font-style:normal;\">\n"
                                      "<p align=\"right\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\">0</p></body></html>"))
        self.time.setText(_translate("MainWindow", "0: 00"))
        self.label.setText(_translate("MainWindow",
                                      "Уважаемые гости! Для безопасности посещения просим вас использовать средства индивидуальной защиты органов дыхания"))
        self.logo_auto.setText(_translate("MainWindow", "Здесь могла бы быть ваша реклама"))
        self.label_2.setText(_translate("MainWindow", "Введется видеонаблюдение, за невыполнение, штраф 5000 рублей"))
        self.start_shopping.setText(_translate("MainWindow", "Начать\n"
                                                             "покупку"))
        self.label_10.setText(_translate("MainWindow", "Маленький пакет"))
        self.label_12.setText(_translate("MainWindow", "Большой пакет"))
        self.label_7.setText(_translate("MainWindow", "Ручной ввод"))


class OrderWindow(object):
    def setupUi(self, MainWindow):
        self.cost = 1
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(800, 791)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.designdown = QtWidgets.QFrame(self.centralwidget)
        self.designdown.setGeometry(QtCore.QRect(0, 730, 801, 21))
        self.designdown.setStyleSheet("background-color: rgb(0, 170, 255);")
        self.designdown.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.designdown.setFrameShadow(QtWidgets.QFrame.Raised)
        self.designdown.setObjectName("designdown")
        self.design_up = QtWidgets.QFrame(self.centralwidget)
        self.design_up.setGeometry(QtCore.QRect(0, 70, 801, 16))
        self.design_up.setStyleSheet("background-color: rgb(0, 170, 255);")
        self.design_up.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.design_up.setFrameShadow(QtWidgets.QFrame.Raised)
        self.design_up.setObjectName("design_up")
        self.time = QtWidgets.QLabel(self.centralwidget)
        self.time.setGeometry(QtCore.QRect(30, 40, 221, 16))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.time.setFont(font)
        self.time.setObjectName("time")
        self.phone = QtWidgets.QLabel(self.centralwidget)
        self.phone.setGeometry(QtCore.QRect(550, 50, 251, 16))
        self.phone.setObjectName("phone")
        self.make_order = QtWidgets.QPushButton(self.centralwidget)
        self.make_order.setGeometry(QtCore.QRect(0, 690, 801, 31))
        self.make_order.setStyleSheet("background-color: rgb(255, 251, 140);")
        self.make_order.setObjectName("make_order")
        self.privacy_policy = QtWidgets.QLabel(self.centralwidget)
        self.privacy_policy.setGeometry(QtCore.QRect(280, 760, 531, 16))
        self.privacy_policy.setObjectName("privacy_policy")
        self.total_amount = QtWidgets.QLabel(self.centralwidget)
        self.total_amount.setGeometry(QtCore.QRect(20, 655, 181, 21))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.total_amount.setFont(font)
        self.total_amount.setObjectName("total_amount")
        self.textBrowser = QtWidgets.QTextBrowser(self.centralwidget)
        self.textBrowser.setGeometry(QtCore.QRect(570, 651, 211, 31))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.textBrowser.setFont(font)
        self.textBrowser.setObjectName("textBrowser")
        self.title_company = QtWidgets.QLabel(self.centralwidget)
        self.title_company.setGeometry(QtCore.QRect(0, 0, 791, 51))
        font = QtGui.QFont()
        font.setFamily("Segoe Script")
        font.setPointSize(14)
        self.title_company.setFont(font)
        self.title_company.setAlignment(QtCore.Qt.AlignCenter)
        self.title_company.setObjectName("title_company")
        self.sale_code = QtWidgets.QPushButton(self.centralwidget)
        self.sale_code.setGeometry(QtCore.QRect(640, 620, 151, 21))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.sale_code.setFont(font)
        self.sale_code.setStyleSheet("border : none;  color: rgb(0, 170, 255);")
        self.sale_code.setIconSize(QtCore.QSize(16, 16))
        self.sale_code.setObjectName("sale_code")
        self.name_card = QtWidgets.QLabel(self.centralwidget)
        self.name_card.setGeometry(QtCore.QRect(180, 240, 131, 21))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.name_card.setFont(font)
        self.name_card.setObjectName("name_card")
        self.name_card_input = QtWidgets.QTextEdit(self.centralwidget)
        self.name_card_input.setGeometry(QtCore.QRect(180, 270, 381, 31))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.name_card_input.setFont(font)
        self.name_card_input.setObjectName("name_card_input")
        self.validity_period = QtWidgets.QLabel(self.centralwidget)
        self.validity_period.setGeometry(QtCore.QRect(180, 320, 131, 21))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.validity_period.setFont(font)
        self.validity_period.setObjectName("validity_period")
        self.cvc = QtWidgets.QLabel(self.centralwidget)
        self.cvc.setGeometry(QtCore.QRect(380, 320, 131, 21))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.cvc.setFont(font)
        self.cvc.setObjectName("cvc")
        self.mail = QtWidgets.QLabel(self.centralwidget)
        self.mail.setGeometry(QtCore.QRect(180, 390, 161, 21))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.mail.setFont(font)
        self.mail.setObjectName("mail")
        self.validiry_period_input = QtWidgets.QTextEdit(self.centralwidget)
        self.validiry_period_input.setGeometry(QtCore.QRect(180, 350, 161, 31))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.validiry_period_input.setFont(font)
        self.validiry_period_input.setObjectName("validiry_period_input")
        self.cvc_input = QtWidgets.QTextEdit(self.centralwidget)
        self.cvc_input.setGeometry(QtCore.QRect(380, 350, 181, 31))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.cvc_input.setFont(font)
        self.cvc_input.setObjectName("cvc_input")
        self.mail_input = QtWidgets.QTextEdit(self.centralwidget)
        self.mail_input.setGeometry(QtCore.QRect(180, 420, 381, 31))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.mail_input.setFont(font)
        self.mail_input.setObjectName("mail_input")
        self.checkBox = QtWidgets.QCheckBox(self.centralwidget)
        self.checkBox.setGeometry(QtCore.QRect(440, 470, 121, 16))
        self.checkBox.setObjectName("checkBox")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(150, 180, 451, 31))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Ignored, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.label.sizePolicy().hasHeightForWidth())
        self.label.setSizePolicy(sizePolicy)
        self.label.setMinimumSize(QtCore.QSize(4, 1))
        font = QtGui.QFont()
        font.setPointSize(15)
        font.setBold(True)
        font.setWeight(75)
        self.label.setFont(font)
        self.label.setFocusPolicy(QtCore.Qt.ClickFocus)
        self.label.setAutoFillBackground(False)
        self.label.setAlignment(QtCore.Qt.AlignCenter)
        self.label.setObjectName("label")
        MainWindow.setCentralWidget(self.centralwidget)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.time.setText(_translate("MainWindow", "0: 00"))
        self.phone.setText(_translate("MainWindow", "телефон поддержки: 7-928-XXX-XX-XX"))
        self.make_order.setText(_translate("MainWindow", "Оплатить"))
        self.privacy_policy.setText(
            _translate("MainWindow", "При покупке вы соглашаетесь с Политикой конфиденциальность нашего магазина"))
        self.total_amount.setText(_translate("MainWindow", "Итоговая сумма"))
        self.textBrowser.setHtml(_translate("MainWindow",
                                            "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
                                            "<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
                                            "p, li { white-space: pre-wrap; }\n"
                                            "</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:12pt; font-weight:400; font-style:normal;\">\n"
                                            "<p align=\"right\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\">0</p></body></html>"))
        self.title_company.setText(_translate("MainWindow", "Здесь могла бы быть ваша реклама"))
        self.sale_code.setText(_translate("MainWindow", "Введите код купона"))
        self.name_card.setText(_translate("MainWindow", "Номер карты"))
        self.name_card_input.setHtml(_translate("MainWindow",
                                                "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
                                                "<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
                                                "p, li { white-space: pre-wrap; }\n"
                                                "</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:12pt; font-weight:400; font-style:normal;\">\n"
                                                "<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\">0000 0000 0000 0000</p>\n"
                                                "<p style=\"-qt-paragraph-type:empty; margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><br /></p></body></html>"))
        self.validity_period.setText(_translate("MainWindow", "Срок действия"))
        self.cvc.setText(_translate("MainWindow", "CVV2/CVC2"))
        self.mail.setText(_translate("MainWindow", "Электронная почта"))
        self.validiry_period_input.setHtml(_translate("MainWindow",
                                                      "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
                                                      "<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
                                                      "p, li { white-space: pre-wrap; }\n"
                                                      "</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:12pt; font-weight:400; font-style:normal;\">\n"
                                                      "<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\">ММ/ГГ</p></body></html>"))
        self.cvc_input.setHtml(_translate("MainWindow",
                                          "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
                                          "<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
                                          "p, li { white-space: pre-wrap; }\n"
                                          "</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:12pt; font-weight:400; font-style:normal;\">\n"
                                          "<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\">***</p></body></html>"))
        self.mail_input.setHtml(_translate("MainWindow",
                                           "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
                                           "<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
                                           "p, li { white-space: pre-wrap; }\n"
                                           "</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:12pt; font-weight:400; font-style:normal;\">\n"
                                           "<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\">aboba@gmail.com</p></body></html>"))
        self.checkBox.setText(_translate("MainWindow", "Распечатать чек?"))
        self.label.setText(_translate("MainWindow", "Указан неверный номер карты"))


class ThxOrderWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(802, 790)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.designdown = QtWidgets.QFrame(self.centralwidget)
        self.designdown.setGeometry(QtCore.QRect(0, 730, 801, 21))
        self.designdown.setStyleSheet("background-color: rgb(0, 170, 255);")
        self.designdown.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.designdown.setFrameShadow(QtWidgets.QFrame.Raised)
        self.designdown.setObjectName("designdown")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(50, 90, 671, 491))
        self.label.setText("")
        self.label.setPixmap(QtGui.QPixmap("icons/thankxxx.jpg"))
        self.label.setScaledContents(True)
        self.label.setObjectName("label")
        self.return_menu = QtWidgets.QPushButton(self.centralwidget)
        self.return_menu.setGeometry(QtCore.QRect(160, 610, 461, 51))
        font = QtGui.QFont()
        font.setPointSize(16)
        self.return_menu.setFont(font)
        self.return_menu.setStyleSheet("background-color: rgb(255, 251, 140);")
        self.return_menu.setObjectName("return_menu")
        self.time = QtWidgets.QLabel(self.centralwidget)
        self.time.setGeometry(QtCore.QRect(30, 50, 211, 16))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.time.setFont(font)
        self.time.setObjectName("time")
        self.title_company = QtWidgets.QLabel(self.centralwidget)
        self.title_company.setGeometry(QtCore.QRect(0, 10, 791, 51))
        font = QtGui.QFont()
        font.setFamily("Segoe Script")
        font.setPointSize(14)
        self.title_company.setFont(font)
        self.title_company.setAlignment(QtCore.Qt.AlignCenter)
        self.title_company.setObjectName("title_company")
        self.phone = QtWidgets.QLabel(self.centralwidget)
        self.phone.setGeometry(QtCore.QRect(550, 60, 261, 16))
        self.phone.setObjectName("phone")
        self.design_up = QtWidgets.QFrame(self.centralwidget)
        self.design_up.setGeometry(QtCore.QRect(0, 80, 801, 16))
        self.design_up.setStyleSheet("background-color: rgb(0, 170, 255);")
        self.design_up.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.design_up.setFrameShadow(QtWidgets.QFrame.Raised)
        self.design_up.setObjectName("design_up")
        MainWindow.setCentralWidget(self.centralwidget)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.return_menu.setText(_translate("MainWindow", "Вернуться в меню"))
        self.time.setText(_translate("MainWindow", "0: 00"))
        self.title_company.setText(_translate("MainWindow", "Здесь могла бы быть ваша реклама"))
        self.phone.setText(_translate("MainWindow", "телефон поддержки: 7-928-XXX-XX-XX"))


class LoginWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(787, 783)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.designdown = QtWidgets.QFrame(self.centralwidget)
        self.designdown.setGeometry(QtCore.QRect(-10, 730, 801, 21))
        self.designdown.setStyleSheet("background-color: rgb(0, 170, 255);")
        self.designdown.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.designdown.setFrameShadow(QtWidgets.QFrame.Raised)
        self.designdown.setObjectName("designdown")
        self.straf = QtWidgets.QLabel(self.centralwidget)
        self.straf.setGeometry(QtCore.QRect(0, 750, 781, 20))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.straf.setFont(font)
        self.straf.setText("")
        self.straf.setAlignment(QtCore.Qt.AlignCenter)
        self.straf.setObjectName("straf")
        self.login = QtWidgets.QLabel(self.centralwidget)
        self.login.setGeometry(QtCore.QRect(220, 160, 161, 31))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setBold(False)
        font.setWeight(50)
        self.login.setFont(font)
        self.login.setObjectName("login")
        self.password = QtWidgets.QLabel(self.centralwidget)
        self.password.setGeometry(QtCore.QRect(220, 200, 161, 31))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setBold(False)
        font.setWeight(50)
        self.password.setFont(font)
        self.password.setObjectName("password")
        self.login_iput = QtWidgets.QTextEdit(self.centralwidget)
        self.login_iput.setGeometry(QtCore.QRect(320, 160, 261, 31))
        self.login_iput.setObjectName("login_iput")
        self.password_input = QtWidgets.QTextEdit(self.centralwidget)
        self.password_input.setGeometry(QtCore.QRect(320, 200, 261, 31))
        self.password_input.setObjectName("password_input")
        self.prank = QtWidgets.QPushButton(self.centralwidget)
        self.prank.setGeometry(QtCore.QRect(470, 240, 111, 31))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.prank.setFont(font)
        self.prank.setAutoFillBackground(False)
        self.prank.setStyleSheet("background-color: rgb(170, 0, 0);")
        self.prank.setObjectName("prank")
        self.title_company = QtWidgets.QLabel(self.centralwidget)
        self.title_company.setGeometry(QtCore.QRect(0, 10, 791, 51))
        font = QtGui.QFont()
        font.setFamily("Segoe Script")
        font.setPointSize(14)
        self.title_company.setFont(font)
        self.title_company.setStyleSheet("color: rgb(255, 0, 4);")
        self.title_company.setAlignment(QtCore.Qt.AlignCenter)
        self.title_company.setObjectName("title_company")
        self.design_up = QtWidgets.QFrame(self.centralwidget)
        self.design_up.setGeometry(QtCore.QRect(0, 70, 801, 16))
        self.design_up.setStyleSheet("background-color: rgb(0, 170, 255);")
        self.design_up.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.design_up.setFrameShadow(QtWidgets.QFrame.Raised)
        self.design_up.setObjectName("design_up")
        self.error = QtWidgets.QLabel(self.centralwidget)
        self.error.setGeometry(QtCore.QRect(220, 250, 241, 21))
        self.error.setText("")
        self.error.setObjectName("error")
        self.autorization = QtWidgets.QPushButton(self.centralwidget)
        self.autorization.setGeometry(QtCore.QRect(240, 750, 321, 23))
        self.autorization.setStyleSheet("border : none;")
        self.autorization.setObjectName("autorization")
        MainWindow.setCentralWidget(self.centralwidget)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.login.setText(_translate("MainWindow", "Логин:"))
        self.password.setText(_translate("MainWindow", "Пароль:"))
        self.prank.setText(_translate("MainWindow", "Войти"))
        self.title_company.setText(_translate("MainWindow", "ЭТО АДМИН МЕНЮ"))
        self.autorization.setText(_translate("MainWindow", "Просим вас не буянить, штраф 100000 рублей"))


class AdminWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(800, 779)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.background = QtWidgets.QLabel(self.centralwidget)
        self.background.setEnabled(True)
        self.background.setGeometry(QtCore.QRect(0, 0, 801, 781))
        self.background.setText("")
        self.background.setPixmap(QtGui.QPixmap("icons/hack.png"))
        self.background.setScaledContents(True)
        self.background.setObjectName("background")
        self.title_company = QtWidgets.QLabel(self.centralwidget)
        self.title_company.setGeometry(QtCore.QRect(0, 50, 791, 51))
        font = QtGui.QFont()
        font.setFamily("Segoe Script")
        font.setPointSize(14)
        self.title_company.setFont(font)
        self.title_company.setStyleSheet("color: rgb(255, 0, 4);")
        self.title_company.setAlignment(QtCore.Qt.AlignCenter)
        self.title_company.setObjectName("title_company")
        self.create_qr = QtWidgets.QPushButton(self.centralwidget)
        self.create_qr.setGeometry(QtCore.QRect(30, 110, 271, 51))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.create_qr.setFont(font)
        self.create_qr.setStyleSheet("border:none;\n"
                                     "color: rgb(85, 0, 127);")
        self.create_qr.setObjectName("create_qr")
        self.return_menu = QtWidgets.QPushButton(self.centralwidget)
        self.return_menu.setGeometry(QtCore.QRect(190, 680, 411, 81))
        font = QtGui.QFont()
        font.setFamily("Segoe Script")
        font.setPointSize(20)
        self.return_menu.setFont(font)
        self.return_menu.setStyleSheet("color: rgb(255, 0, 4); border : none;")
        self.return_menu.setObjectName("return_menu")
        self.change_password = QtWidgets.QPushButton(self.centralwidget)
        self.change_password.setGeometry(QtCore.QRect(290, 170, 211, 51))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.change_password.setFont(font)
        self.change_password.setStyleSheet("border:none;\n"
                                           "color: rgb(85, 0, 127);")
        self.change_password.setObjectName("change_password")
        self.list_products = QtWidgets.QPushButton(self.centralwidget)
        self.list_products.setGeometry(QtCore.QRect(460, 110, 261, 51))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.list_products.setFont(font)
        self.list_products.setStyleSheet("border:none;\n"
                                         "color: rgb(85, 0, 127);")
        self.list_products.setObjectName("list_products")
        self.change_password_2 = QtWidgets.QPushButton(self.centralwidget)
        self.change_password_2.setGeometry(QtCore.QRect(290, 450, 211, 51))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.change_password_2.setFont(font)
        self.change_password_2.setStyleSheet("border:none;\n"
                                             "color: rgb(85, 0, 127);")
        self.change_password_2.setObjectName("change_password_2")
        MainWindow.setCentralWidget(self.centralwidget)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.title_company.setText(_translate("MainWindow", "ЭТО АДМИН МЕНЮ"))
        self.create_qr.setText(_translate("MainWindow", "Cоздать QR код"))
        self.return_menu.setText(_translate("MainWindow", "Вернуться обратно"))
        self.change_password.setText(_translate("MainWindow", "Изменить пароль"))
        self.list_products.setText(_translate("MainWindow", "Просмотреть список товаров"))
        self.change_password_2.setText(_translate("MainWindow", "Добавить купон"))


class TableWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(800, 780)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.tableWidget = QtWidgets.QTableWidget(self.centralwidget)
        self.tableWidget.setGeometry(QtCore.QRect(0, 50, 800, 720))
        self.tableWidget.setObjectName("tableWidget")
        self.tableWidget.setColumnCount(0)
        self.tableWidget.setRowCount(0)
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(10, 10, 201, 21))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label.setFont(font)
        self.label.setObjectName("label")
        self.input_words = QtWidgets.QLineEdit(self.centralwidget)
        self.input_words.setGeometry(QtCore.QRect(190, 10, 581, 21))
        self.input_words.setObjectName("input_words")
        MainWindow.setCentralWidget(self.centralwidget)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.label.setText(_translate("MainWindow", "Подстрока для покупателя"))


class Open_MainWindow(QMainWindow, MainWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.time.setText(str(datetime.datetime.now())[:10])
        self.create_table()
        self.start_shopping.clicked.connect(self.open_order)
        self.logo_auto.clicked.connect(self.open_login)
        self.login_window = OpenLoginWindow()
        self.small_package.clicked.connect(partial(self.add_item, self.small_package.objectName()))
        self.big_package.clicked.connect(partial(self.add_item, self.big_package.objectName()))
        self.col_items.itemChanged.connect(self.change_price)
        self.flag = False
        self.manual_input.clicked.connect(self.manual)
        self.drawf()

    def manual(self):
        """Если не получается отсканиовать товар,
        то можно вручную ввести код товара"""
        index, ok_pressed = QInputDialog.getText(self, "Обновление",
                                                 "Введите индекс товара:")
        if ok_pressed:
            self.col_items.setRowCount(
                self.col_items.rowCount() + 1)
            global con
            cur = con.cursor()
            """загрузить товар с нужным индексом"""
            products = cur.execute(f"""SELECT * FROM 'products'
            WHERE ind = ?""", (index,)).fetchone()
            for i, row in enumerate(products[1:3]):
                print(i, row)
                item = QTableWidgetItem(str(row))
                item.setFlags(QtCore.Qt.ItemIsEnabled)
                self.col_items.setItem(
                    self.col_items.rowCount() - 1, i, item)
            self.col_items.setItem(
                self.col_items.rowCount() - 1, 2, QTableWidgetItem('1'))
            self.col_items.resizeColumnsToContents()

    def drawf(self):
        """Начало отрисовки логотипа"""
        self.flag = True
        self.update()

    def paintEvent(self, event):
        if self.flag:
            self.qp = QPainter()
            self.qp.begin(self)
            self.qp.setPen(Qt.blue)
            self.drawSq()
            self.qp.end()

    def scale(self, p1, p2, k):
        return p1[0] + k * (p2[0] - p1[0]), p1[1] + k * (p2[1] - p1[1])

    def create_polygon(self, nodes):
        nodes2 = [(int(node[0] + 1400 // 2), int(200 // 2 - node[1])) for node in nodes]
        for i in range(-1, len(nodes2) - 1):
            self.qp.drawLine(*nodes2[i], *nodes2[i + 1])

    def drawSq(self):
        """Рисует логотип компании"""
        kel = float(0.9)
        p = int(8)
        RAD = 80
        nodes = [(RAD * cos(i * 2 * pi / p), RAD * sin(i * 2 * pi / p)) for i in range(p)]
        for i in range(int(10)):
            self.create_polygon(nodes)
            new_nodes = []
            for index in range(-1, len(nodes) - 1):
                new_nodes.append(self.scale(nodes[index], nodes[index + 1], kel))
            nodes = new_nodes[:]

    def change_price(self):
        """Считывает изменения таблицы, для подсчета цены"""
        rows = self.col_items.rowCount()
        price = 0
        for row in range(rows):
            try:
                price += (int(self.col_items.item(row, 2).text()) * int(self.col_items.item(row, 1).text()))
                if self.col_items.item(row, 2).text() == '0':
                    """Если кол-во товара выставленно нулем, 
                    перекасить в серый цвет(сделать неактивным)"""
                    for i in range(3):
                        self.col_items.item(row, i).setBackground(QtGui.QBrush(QtGui.QColor("#C0C0C0")))
                else:
                    for i in range(3):
                        self.col_items.item(row, i).setBackground(QtGui.QBrush(QtGui.QColor("#FFFFFF")))
            except:
                """Если кол-во еще не обновилось, пропустить"""
                pass
            self.money.setText(str(price))

    def create_table(self):
        """Создать или перезагрузить таблицу"""
        self.col_items.setColumnCount(3)
        self.col_items.setHorizontalHeaderLabels(["Название", "Цена", "Количество"])
        self.col_items.setRowCount(0)
        self.money.setText('0')

    def open_login(self):
        """Секретное окно со входом в админ меню"""
        self.login_window.show()

    def add_item(self, name):
        """Обработка допольнительных кнопок для добавления пакетов
        input:
            str(name) :
                    1) big_package
                    2) small_package"""
        self.col_items.setRowCount(
            self.col_items.rowCount() + 1)
        if name == 'big_package':
            price = QTableWidgetItem(str(30))
        else:
            price = QTableWidgetItem(str(15))
        name = QTableWidgetItem(str(name))
        name.setFlags(QtCore.Qt.ItemIsEnabled)
        price.setFlags(QtCore.Qt.ItemIsEnabled)
        self.col_items.setItem(
            self.col_items.rowCount() - 1, 0, name)
        self.col_items.setItem(
            self.col_items.rowCount() - 1, 1, price)
        self.col_items.setItem(
            self.col_items.rowCount() - 1, 2, QTableWidgetItem("1"))
        self.col_items.resizeColumnsToContents()

    def open_order(self):
        """Октрыть окно с оплатой товаров"""
        check_info = []
        for i in range(self.col_items.rowCount()):
            for j in range(self.col_items.columnCount()):
                check_info.append(str(self.col_items.item(i, j).text()))
        self.to_order = Open_OrderWindow(str(self.money.toPlainText()), check_info)
        self.create_table()
        self.to_order.show()


class Open_OrderWindow(QMainWindow, OrderWindow):
    def __init__(self, price, check_info):
        super().__init__()
        self.setupUi(self)
        self.label.setStyleSheet("color: rgb(255, 0, 4);")
        self.time.setText(str(datetime.datetime.now())[:10])
        self.label.hide()
        self.price = price
        self.sale_code.clicked.connect(self.run)
        self.check_info = check_info
        self.to_ThxOrder = OpenThxOrderWindow()
        self.make_order.clicked.connect(self.open_ThxOrder)
        self.textBrowser.setText(price)
        self.code = ''
        self.flag = False
        self.drawf()

    def drawf(self):
        """Начало отрисовки логотипа"""
        self.flag = True
        self.update()

    def paintEvent(self, event):
        if self.flag:
            self.qp = QPainter()
            self.qp.begin(self)
            self.qp.setPen(Qt.blue)
            self.drawSq()
            self.qp.end()

    def scale(self, p1, p2, k):
        return p1[0] + k * (p2[0] - p1[0]), p1[1] + k * (p2[1] - p1[1])

    def create_polygon(self, nodes):
        nodes2 = [(int(node[0] + 1400 // 2), int(200 // 2 - node[1])) for node in nodes]
        for i in range(-1, len(nodes2) - 1):
            self.qp.drawLine(*nodes2[i], *nodes2[i + 1])

    def drawSq(self):
        """Рисует логотип компании"""
        kel = float(0.9)
        p = int(8)
        RAD = 80
        nodes = [(RAD * cos(i * 2 * pi / p), RAD * sin(i * 2 * pi / p)) for i in range(p)]
        for i in range(int(10)):
            self.create_polygon(nodes)
            new_nodes = []
            for index in range(-1, len(nodes) - 1):
                new_nodes.append(self.scale(nodes[index], nodes[index + 1], kel))
            nodes = new_nodes[:]

    def run(self):
        self.code, ok_pressed = QInputDialog.getText(self, "Введите ваш промокод", "Код")
        codes = open('promo.txt', encoding='utf8').readlines()
        self.cost = 1
        for i in codes:
            if i.split()[0] == self.code:
                self.cost = 1 - int(i.split()[1]) / 100
                self.textBrowser.setText(str(int(self.price) * float(self.cost)))
                break

    def check(self):
        document = docx.Document()
        document.add_heading('НеКвадрат', 0)
        document.add_paragraph('--------------------------------'
                               '----------------------------------'
                               '---------------------------------'
                               '-------------------')
        table = document.add_table(rows=1 + len(self.check_info) // 3, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Название'
        hdr_cells[1].text = 'Цена'
        hdr_cells[2].text = 'Количество'
        hdr_cells[3].text = 'Итого'
        sum_ = 0
        for i in range(1, len(self.check_info) // 3 + 1):
            hdr_cells = table.rows[i].cells
            for j in range(3):
                hdr_cells[j].text = str(self.check_info[(i - 1) * 3 + j])
            hdr_cells[3].text = str(int(hdr_cells[1].text) * int(hdr_cells[2].text))
            sum_ += int(hdr_cells[3].text)
        document.add_paragraph('--------------------------------'
                               '----------------------------------'
                               '---------------------------------'
                               '-------------------\n')
        document.add_paragraph(f'Итого к оплате(c учетом скидки): {sum_ * self.cost}')
        document.add_paragraph(f'Время покупки: {str(datetime.datetime.now())[:19]}')
        document.add_page_break()

        document.save(f'checks/Check {"-".join(str(datetime.datetime.now())[:19].split(":"))}.docx')

    def luhn(self, code):
        """Алгоритм Луна, проверяет верность номера банковской карты"""
        LOOKUP = (0, 2, 4, 6, 8, 1, 3, 5, 7, 9)
        code = reduce(str.__add__, filter(str.isdigit, code))
        evens = sum(int(i) for i in code[-1::-2])
        odds = sum(LOOKUP[int(i)] for i in code[-2::-2])
        return (evens + odds) % 10 == 0

    def open_ThxOrder(self):
        if self.luhn(self.name_card_input.toPlainText()) and len(
                ''.join(self.name_card_input.toPlainText().split())) == 16:
            """Октрывает окно 'Спасибо за покупку'
            Закрывает окно с заказом"""
            self.check()
            self.to_ThxOrder.show()
            self.hide()
        else:
            self.label.show()


class OpenThxOrderWindow(QMainWindow, ThxOrderWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.time.setText(str(datetime.datetime.now())[:10])
        self.return_menu.clicked.connect(self.open_main)

    def open_main(self):
        """Закрывает окно 'Спасибо за покупку'"""
        self.hide()


class OpenLoginWindow(QMainWindow, LoginWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.autorization.clicked.connect(self.open_admin)
        self.admin_window = OpenAdminWindow()
        self.login_iput.setText("admin")
        self.password_input.setText("admin")
        self.prank.clicked.connect(self.pranker)

    def pranker(self):
        """Кнопка которая убегает, если на нее нажать
        Пранк. :)"""
        self.prank.move(random.randint(0, 800), random.randint(0, 600))

    def open_admin(self):
        if login == self.login_iput.toPlainText() and current_password == self.password_input.toPlainText():
            """Вход в админ меню
            Закрытие логин меню"""
            self.admin_window.show()
            self.login_iput.setText("")
            self.password_input.setText("")
            self.hide()
        else:
            """Если пароль неверный забоговать программу"""
            self.error.setText("Пароль неверный с вас 100 тысяч...")
            self.centralwidget.setStyleSheet("background-color: rgb(170, 0, 0);")
            self.login_iput.setEnabled(False)
            self.password_input.setEnabled(False)


class OpenAdminWindow(QMainWindow, AdminWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.return_menu.clicked.connect(self.open_main)
        self.change_password.clicked.connect(self.change)
        self.create_qr.clicked.connect(self.create)
        self.table_window = OpenTableWindow()
        self.list_products.clicked.connect(self.open_table)
        self.change_password_2.clicked.connect(self.add_code)

    def add_code(self):
        promo, ok_pressed = QInputDialog.getText(self, "Добавление промокода",
                                                 "Введите новый пробел и скидку"
                                                 " через пробел. Например: Test 1")
        codes = open('promo.txt', encoding='utf8').readlines()
        codes.append(promo)
        file = open('promo.txt', 'wt', encoding='utf8')
        file.write('\n'.join(codes))
        file.close()

    def open_table(self):
        """Открыть окно с таблицей товаров"""
        self.table_window.show()

    def change(self):
        """Поменять пароль для входа в админ меню"""
        name, ok_pressed = QInputDialog.getText(self, "Смена пароля.",
                                                "Введите пароль:")
        if ok_pressed:
            global current_password
            current_password = str(name)

    def create(self):
        """Создать bar_code по индексу товара"""

        name, ok_pressed = QInputDialog.getText(self, "Создание qr-кодов.",
                                                "Введите индекс:")
        if ok_pressed and len(name) == 13:
            my_code = barcode.EAN13(name)
            my_code.save("codes/new_code")
        else:
            dialog = QDialog()
            btn = QLabel('ERROR', dialog)
            btn.setStyleSheet("color: rgb(255, 0, 4);")
            btn.move(30, 10)
            dialog.setWindowTitle("Dialog")
            dialog.setWindowModality(Qt.ApplicationModal)
            dialog.exec_()

    def open_main(self):
        self.hide()


class OpenTableWindow(QMainWindow, TableWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.create_table()
        self.input_words.textChanged.connect(self.change)

    def change(self):
        """Форматированный вывод таблицы"""
        if len(self.input_words.text()) > 2:
            self.create_table(self.input_words.text().lower())
        else:
            self.create_table()

    def create_table(self, name=""):
        """Выводит данные в таблицу"""
        global con
        cur = con.cursor()
        products = cur.execute("""SELECT * FROM 'products'""").fetchall()
        self.tableWidget.setColumnCount(4)
        self.tableWidget.setHorizontalHeaderLabels(["Индекс", "Название", "Цена", "Всего продукции"])
        self.tableWidget.setRowCount(0)
        col = 0
        for i, row in enumerate(products):
            if len(row) == 0 or name not in row[1].lower():
                continue
            self.tableWidget.setRowCount(
                self.tableWidget.rowCount() + 1)
            for j, elem in enumerate(row):
                item = QTableWidgetItem(str(elem))
                self.tableWidget.setItem(
                    col, j, item)
            index = 0
            alp = """abcdefghijklmnopqrstuvwxyz
            ABCDEFGHIJKLMNOPQRSTUVWXYZ
            абвгдеёжзийклмнопрстуфхцчшщъыьэюя
            АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ"""
            name_pr = self.tableWidget.item(col, 1).text()
            for symbol in name_pr:
                if symbol in alp:
                    index += alp.index(symbol)
                else:
                    index += 10
            while len(str(index)) != 13:
                index *= 10 + 1
            self.tableWidget.setItem(
                col, 0, QTableWidgetItem(str(index)))
            col += 1
            print(name_pr, ':', index)
            cur.execute("""UPDATE 'products' SET
            ind = ? 
                WHERE name_products = ?""", (str(index), name_pr))
        cur.close()
        self.tableWidget.resizeColumnsToContents()


con = sqlite3.connect("categori.db")

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = Open_MainWindow()
    ex.show()
    sys.exit(app.exec_())
